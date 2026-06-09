# services/ocr.py
"""
OCR 双引擎 + 药品名称提取
===========================
提供 macOS Vision → Tesseract 降级 OCR 和 LangChain 药品名称提取。
"""

import io
import os
import uuid
import zipfile

from ocrmac import ocrmac
import pytesseract
from PIL import Image

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from core.logging_config import logger
from core.config import llm, TEMP_DIR

PDF_DIRECT_TEXT_THRESHOLD = 80
PDF_MIN_OCR_TEXT_LENGTH = 20
PDF_MAX_RENDER_SIDE = 2200
PDF_MIN_IMAGE_AREA = 120_000
PDF_MAX_IMAGES_PER_PAGE = 2
DOCX_MIN_TEXT_THRESHOLD = 50
DOCX_MAX_IMAGES_FOR_OCR = 3

# 药品名称提取链
_name_extraction_prompt = ChatPromptTemplate.from_messages([
    ("system", """你是一个专业药剂师。请从OCR文本中提取**药品名称**。
    规则：
    1. 优先提取通用名（如阿莫西林胶囊）或商品名。
    2. 如果找不到明确药名，请提取**说明书标题**或**第一行显著的文字**。
    3. **必须输出中文**。如果识别到的是英文名，请尝试翻译成中文通用名。
    4. 只输出名称本身，不要包含"药品名称:"等前缀。
    5. 只有在完全无法识别任何有意义文本时，才输出"未知药品"。"""),
    ("user", "{text}")
])
_name_chain = _name_extraction_prompt | llm | StrOutputParser()


def run_ocr(image_path: str, min_text_length: int = 20) -> str:
    """图片 OCR 识别（双引擎降级）
    优先使用 macOS Vision (ocrmac)，识别文字过短或失败时降级到 Tesseract。
    返回识别出的文本，取两个引擎中结果更长的。"""
    vision_text = ""

    try:
        ocr_res = ocrmac.OCR(image_path, language_preference=["zh-Hans"]).recognize()
        vision_text = "\n".join([item[0] for item in ocr_res]).strip()
        logger.info(f" ocrmac produced {len(vision_text)} chars: {vision_text[:200]}")
        if len(vision_text) >= min_text_length:
            logger.info(f" OCR Engine: macOS Vision (ocrmac) — {len(vision_text)} chars")
            return vision_text
    except Exception as e:
        logger.warning(f" ocrmac failed: {e}")

    # Fallback: Tesseract OCR
    if vision_text:
        logger.info(
            f" OCR Engine: macOS Vision text too short ({len(vision_text)} chars), fallback to Tesseract"
        )
    logger.info(f" OCR Engine: Tesseract (fallback)")
    with Image.open(image_path) as img:
        tesseract_text = pytesseract.image_to_string(img, lang="chi_sim+eng").strip()
    logger.info(f" Tesseract produced {len(tesseract_text)} chars: {tesseract_text[:200]}")

    if len(tesseract_text) > len(vision_text):
        return tesseract_text
    return vision_text or tesseract_text


def _write_temp_image_bytes(image_bytes: bytes, suffix: str, prefix: str) -> str:
    """将图片字节流写入临时文件（OCR引擎需要文件路径，不支持直接读字节流）"""
    temp_name = f"{prefix}_{uuid.uuid4().hex}{suffix}"
    temp_path = os.path.join(TEMP_DIR, temp_name)
    with open(temp_path, "wb") as temp_file:
        temp_file.write(image_bytes)
    return temp_path


def _cleanup_temp_file(temp_path: str) -> None:
    """删除临时文件，OCR完成后清理磁盘"""
    if os.path.exists(temp_path):
        os.remove(temp_path)


def _save_pdf_xref_to_png(doc, xref: int, prefix: str) -> str:
    """将PDF中嵌入的图片对象（通过xref引用）导出为PNG临时文件"""
    import fitz

    temp_path = os.path.join(TEMP_DIR, f"{prefix}_{uuid.uuid4().hex}.png")
    source_pix = fitz.Pixmap(doc, xref)
    pix = source_pix

    if source_pix.alpha or source_pix.n not in (1, 3):
        pix = fitz.Pixmap(fitz.csRGB, source_pix)

    pix.save(temp_path)
    return temp_path


def _collect_pdf_image_candidates(page) -> list[dict]:
    """收集PDF页面中面积足够大的图片，按覆盖率和像素面积降序排列
    过滤掉小图（<120000像素），优先OCR占页面比例大的图片（说明书主要内容通常满页）"""
    page_area = max(page.rect.width * page.rect.height, 1)
    seen_xrefs = set()
    candidates = []

    for image_info in page.get_images(full=True):
        xref = image_info[0]
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)

        width = int(image_info[2] or 0)
        height = int(image_info[3] or 0)
        pixel_area = width * height
        if pixel_area < PDF_MIN_IMAGE_AREA:
            continue

        try:
            rects = page.get_image_rects(xref)
        except Exception:
            rects = []

        coverage = 0.0
        if rects:
            coverage = max((rect.width * rect.height) / page_area for rect in rects)

        candidates.append({
            "xref": xref,
            "pixel_area": pixel_area,
            "coverage": coverage,
        })

    candidates.sort(key=lambda item: (item["coverage"], item["pixel_area"]), reverse=True)
    return candidates


def _ocr_pdf_page_images(doc, page, page_num: int) -> str:
    """对PDF页面中嵌入的图片逐张做OCR，最多处理2张大图
    一旦某张图识别出足够文字就提前返回，避免浪费时间"""
    texts = []
    candidates = _collect_pdf_image_candidates(page)

    for candidate in candidates[:PDF_MAX_IMAGES_PER_PAGE]:
        temp_path = _save_pdf_xref_to_png(doc, candidate["xref"], prefix=f"pdf_page_{page_num}_img")
        try:
            img_text = run_ocr(temp_path, min_text_length=PDF_MIN_OCR_TEXT_LENGTH).strip()
        finally:
            _cleanup_temp_file(temp_path)

        if not img_text:
            continue

        texts.append(img_text)

        if len(img_text) >= PDF_MIN_OCR_TEXT_LENGTH or candidate["coverage"] >= 0.6:
            break

    return "\n".join(texts).strip()


def _ocr_rendered_pdf_page(page, page_num: int) -> str:
    """PDF整页渲染OCR（兜底方案）
    将整页渲染成灰度PNG图片，再调用run_ocr识别，用于纯图片型PDF"""
    import fitz

    longest_side = max(page.rect.width, page.rect.height, 1)
    zoom = min(PDF_MAX_RENDER_SIDE / longest_side, 3.0)
    zoom = max(zoom, 2.0)

    temp_path = os.path.join(TEMP_DIR, f"pdf_render_{page_num}_{uuid.uuid4().hex}.png")
    pix = page.get_pixmap(
        matrix=fitz.Matrix(zoom, zoom),
        colorspace=fitz.csGRAY,
        alpha=False,
    )
    pix.save(temp_path)

    try:
        return run_ocr(temp_path, min_text_length=PDF_MIN_OCR_TEXT_LENGTH).strip()
    finally:
        _cleanup_temp_file(temp_path)


def extract_text_from_pdf(pdf_path: str) -> str:
    """PDF文本提取主入口（三层降级）
    ① 直接提取原生文本（文字型PDF，最快）
    ② 提取嵌入图片做OCR（扫描型PDF）
    ③ 整页渲染成图片再OCR（纯图片PDF，兜底）
    用PyMuPDF (fitz)库操作PDF。"""
    import fitz

    parts = []
    with fitz.open(pdf_path) as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text").strip()
            if page_text:
                parts.append(page_text)

            if len(page_text) >= PDF_DIRECT_TEXT_THRESHOLD:
                continue

            logger.info(f" PDF page {page_num}: trying embedded image OCR first")
            image_text = _ocr_pdf_page_images(doc, page, page_num)
            if image_text:
                parts.append(f"[Page {page_num} OCR]\n{image_text}")
                continue

            logger.info(f" PDF page {page_num}: fallback to page rendering OCR")
            rendered_text = _ocr_rendered_pdf_page(page, page_num)
            if rendered_text:
                parts.append(f"[Page {page_num} OCR]\n{rendered_text}")

    return "\n".join(parts).strip()


def extract_text_from_docx(docx_path: str) -> str:
    """Word文档文本提取（两层降级）
    ① 直接读段落+表格文本
    ② 文字太少（<50字）时，解压docx（本质是zip），对word/media/下的大图做OCR
    最多处理3张最大的图片。"""
    import docx

    parts = []
    doc = docx.Document(docx_path)

    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text:
                        parts.append(para.text)

    raw_text = "\n".join(parts).strip()
    if len(raw_text) >= DOCX_MIN_TEXT_THRESHOLD:
        return raw_text

    logger.info(" Little text found in Docx. Checking for large embedded images...")
    image_candidates = []

    with zipfile.ZipFile(docx_path, "r") as archive:
        for img_name in archive.namelist():
            if not img_name.startswith("word/media/"):
                continue

            image_bytes = archive.read(img_name)
            try:
                with Image.open(io.BytesIO(image_bytes)) as image:
                    width, height = image.size
            except Exception:
                logger.warning(f" Skipping unreadable Docx image: {img_name}")
                continue

            if width * height < PDF_MIN_IMAGE_AREA:
                continue

            image_candidates.append((width * height, img_name, image_bytes))

    image_candidates.sort(key=lambda item: item[0], reverse=True)
    logger.info(f" Found {len(image_candidates)} large images in Docx")

    for _, img_name, image_bytes in image_candidates[:DOCX_MAX_IMAGES_FOR_OCR]:
        suffix = os.path.splitext(img_name)[1] or ".png"
        temp_path = _write_temp_image_bytes(image_bytes, suffix=suffix, prefix="docx_media")
        try:
            img_text = run_ocr(temp_path, min_text_length=PDF_MIN_OCR_TEXT_LENGTH).strip()
        finally:
            _cleanup_temp_file(temp_path)

        if img_text:
            parts.append(f"[Image: {img_name}]\n{img_text}")

    return "\n".join(parts).strip()


def extract_medicine_name(ocr_text: str) -> str:
    """从OCR文本中用LLM提取药品名称
    将OCR文本（截取前1500字）交给DeepSeek，prompt要求优先提取通用名/商品名，
    必须输出中文，只返回名称本身。文本太短（<5字）时直接返回"未知药品"。"""
    logger.info(f" Raw OCR Text for Name Extraction: {ocr_text[:200]}...")

    # 空文本/短文本防御
    if not ocr_text or len(ocr_text.strip()) < 5:
        logger.warning(" Input text too short. Skipping LLM extraction.")
        return "未知药品"

    try:
        name = _name_chain.invoke({"text": ocr_text[:1500]})
        name = name.strip().replace("药品名称：", "").replace("名称：", "")
        if len(name) > 20:
            name = name[:20] + "..."
        return name
    except Exception as e:
        logger.error(f"LangChain Extraction Error: {e}")
        return "未知药品"
